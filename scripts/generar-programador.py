# -*- coding: utf-8 -*-
"""
Genera el Programador de «Investigación en Cuidados Paliativos y Rehabilitación»
rellenando la plantilla institucional «4. Formato. Programador.docx».

Se rellena la plantilla real, no se recrea: así el documento conserva los
estilos, márgenes y encabezados de la Vicerrectoría Académica.

Fuentes:
  · Contenido:  docs/syllabus/Investigación en Cuidados Paliativos y Rehabilitación.docx
  · Calendario: docs/HORARIOS 2026 - II.pdf  (Dirección de la Maestría)
  · Convención de diligenciamiento: programador del Seminario de fundamentación 2026-II.

Requiere:  pip install python-docx
Uso:       python scripts/generar-programador.py
"""
import copy
import sys
import pathlib
import docx
from docx.shared import Pt

RAIZ = pathlib.Path(__file__).resolve().parent.parent
PLANTILLA = RAIZ / 'docs' / '4. Formato. Programador.docx'
SALIDA = RAIZ / 'docs' / 'Programador_2026-II.docx'

DOCENTE = 'Jorge Wilhem Bogoya López'

# ── Datos de identificación ───────────────────────────────────────────
# Las horas se toman del contenido programático oficial (144/48/96). El
# espacio es teórico, así que las teóricas igualan a las presenciales
# declaradas y no hay prácticas. Ver la nota al pie del documento.
DATOS = {
    1: ('Programa: \nMaestría en Cuidados Paliativos',
        'Asignatura:\nInvestigación en Cuidados Paliativos y Rehabilitación'),
    2: ('Código:  \n37543013', 'Plan de estudios:\n2377'),
    3: ('Número de Créditos dentro del Plan de Estudios: 3',
        'Fecha de actualización:\nAgosto de 2026'),
    4: ('Horas Totales', '144'),
    5: ('Horas presenciales', '48'),
    6: ('Horas de trabajo independiente', '96'),
    7: ('Horas teóricas', '48'),
    8: ('Horas prácticas', '0'),
}

# ── Contenidos ────────────────────────────────────────────────────────
# Un renglón por bloque horario, como en el programador del Seminario.
# Franja de este espacio académico: VIERNES de 2:00 a 5:00 p. m., a partir
# del 21 de agosto de 2026 (HORARIOS 2026-II). Siete sesiones. La modalidad
# —presencial o sincrónica— es la que fija el calendario para cada fin de
# semana. Semanas contadas desde la inducción del 31 de julio (semana 1).
# (semana, fecha, bloque, tema, responsable)
FILAS = [
    # ── Sesión 1 · vie 21 ago · Presencial · Presentación + Unidad 1 ──
    ('4', '21/08/2026', '2:00-3:00 pm',
     'Encuentro presencial\n'
     'Presentación del espacio académico y del hilo del curso\n'
     'Unidad 1. Abordajes de investigaciones en rehabilitación paliativa\n'
     'Tema 1. Declinación funcional, trayectorias de enfermedad y concepto de '
     'rehabilitación paliativa', DOCENTE),
    ('', '21/08/2026', '3:00-4:00 pm',
     'Tema 2. Modelos de intervenciones en rehabilitación paliativa', ''),
    ('', '21/08/2026', '4:00-5:00 pm',
     'Tema 3. Realidad virtual, tecnología móvil, aplicaciones en salud, objetos '
     'virtuales de aprendizaje e inteligencia artificial generativa (Wysa)', ''),

    # ── Sesión 2 · vie 4 sep · Sincrónico · Unidad 2 · cierre C1 ──────
    ('6', '04/09/2026', '2:00-3:00 pm',
     'Encuentro sincrónico\n'
     'Unidad 2. Abordajes de investigaciones en rehabilitación paliativa en cáncer\n'
     'Tema 1. Rehabilitación física y psicosocial: dolor, complicaciones neuromusculares '
     'y musculoesqueléticas, edema y linfedema', DOCENTE),
    ('', '04/09/2026', '3:00-4:00 pm',
     'Tema 2. Ámbitos de rehabilitación: unidades paliativas y entornos hospitalarios, '
     'centros día, programas comunitarios y rehabilitación infantil', ''),
    ('', '04/09/2026', '4:00-5:00 pm',
     'Tema 3. Medición de resultados funcionales\n'
     'Entrega Primer Producto 35 %', ''),

    # ── Sesión 3 · vie 18 sep · Presencial · Unidad 3 ────────────────
    ('8', '18/09/2026', '2:00-3:00 pm',
     'Encuentro presencial\n'
     'Unidad 3. Abordajes de investigaciones en rehabilitación paliativa en falla de órgano\n'
     'Tema 1. Rehabilitación física, funcional y social: rehabilitación cardiovascular '
     'e indicaciones para el control de síntomas', DOCENTE),
    ('', '18/09/2026', '3:00-4:00 pm',
     'Tema 1. Manejo de la disfagia orofaríngea y trastornos de la deglución\n'
     'Tema 2. Ámbitos de rehabilitación (institucional, domiciliaria, comunitaria)', ''),
    ('', '18/09/2026', '4:00-5:00 pm',
     'Tema 2. Programas de rehabilitación cardiopulmonar y rehabilitación domiciliaria\n'
     'Tema 3. Medición de resultados funcionales', ''),

    # ── Sesión 4 · vie 2 oct · Sincrónico · Unidad 4 (en inglés) ─────
    ('10', '02/10/2026', '2:00-3:00 pm',
     'Encuentro sincrónico (unidad desarrollada en inglés)\n'
     'Unidad 4. Abordajes de investigaciones en rehabilitación paliativa en trastornos '
     'neurodegenerativos\n'
     'Tema 1. Rehabilitación en esclerosis múltiple y esclerosis lateral amiotrófica', DOCENTE),
    ('', '02/10/2026', '3:00-4:00 pm',
     'Tema 1. Parálisis cerebral\n'
     'Tema 2. Ámbitos de rehabilitación: unidades de larga duración y programas de '
     'rehabilitación integral', ''),
    ('', '02/10/2026', '4:00-5:00 pm',
     'Tema 3. Medición de resultados funcionales en enfermedad neurodegenerativa', ''),

    # ── Sesión 5 · vie 16 oct · Presencial · Unidad 5 · cierre C2 ────
    ('12', '16/10/2026', '2:00-3:00 pm',
     'Encuentro presencial\n'
     'Unidad 5. Abordajes de investigaciones en rehabilitación paliativa en demencia y '
     'fragilidad\n'
     'Tema 1. Rehabilitación física y psicosocial: rehabilitación en demencias', DOCENTE),
    ('', '16/10/2026', '3:00-4:00 pm',
     'Tema 2. Ámbitos de rehabilitación: unidades de larga duración y programas de '
     'rehabilitación integral', ''),
    ('', '16/10/2026', '4:00-5:00 pm',
     'Tema 3. Medición de resultados en demencia y fragilidad: cambio mínimo importante\n'
     'Entrega Segundo Producto 35 %', ''),

    # ── Sesión 6 · vie 30 oct · Sincrónico · Unidad 6 ────────────────
    ('14', '30/10/2026', '2:00-3:00 pm',
     'Encuentro sincrónico\n'
     'Unidad 6. Abordajes de investigaciones en rehabilitación paliativa en eventos '
     'catastróficos (ictus severo, fractura de cadera del adulto mayor)\n'
     'Tema 1. Rehabilitación física y psicosocial: rehabilitación en geriatría y espasticidad', DOCENTE),
    ('', '30/10/2026', '3:00-4:00 pm',
     'Tema 2. Ámbitos de rehabilitación (institucional, domiciliaria, comunitaria)', ''),
    ('', '30/10/2026', '4:00-5:00 pm',
     'Tema 3. Medición de resultados en rehabilitación del adulto mayor', ''),

    # ── Encuentro institucional, ajeno a este espacio académico ──────
    ('16', '13/11/2026', '',
     'Encuentro de investigación — sustentaciones de tesis\n'
     '(Solo para estudiantes de Tesis II. No corresponde a este espacio académico)', ''),

    # ── Sesión 7 · vie 20 nov · Presencial · integrador · cierre C3 ──
    ('17', '20/11/2026', '2:00-3:00 pm',
     'Encuentro presencial\n'
     'Seminario integrador de las seis unidades del espacio académico', DOCENTE),
    ('', '20/11/2026', '3:00-4:00 pm',
     'Socialización de los productos finales de investigación en rehabilitación paliativa', ''),
    ('', '20/11/2026', '4:00-5:00 pm',
     'Coevaluación, autoevaluación y cierre del espacio académico\n'
     'Entrega Producto Final 30 %', ''),
]

NOTA = (
    'Nota. La programación sigue el calendario oficial de encuentros académicos 2026-II de la '
    'Maestría (comunicación de la Dirección, «HORARIOS 2026 - II») y la franja de este espacio '
    'académico, viernes de 2:00 a 5:00 p. m. a partir del 21 de agosto de 2026, con siete sesiones '
    '(21 de agosto; 4 y 18 de septiembre; 2, 16 y 30 de octubre; 20 de noviembre). La modalidad de '
    'cada sesión —presencial o sincrónica— corresponde a la fijada por el calendario para ese fin '
    'de semana: presencial el 21 de agosto, 18 de septiembre, 16 de octubre y 20 de noviembre; '
    'sincrónica (asistida por tecnología) el 4 de septiembre, 2 y 30 de octubre. El encuentro del '
    '13 de noviembre es de sustentaciones de tesis II y no corresponde a este espacio académico. '
    'Las 48 horas presenciales declaradas en el contenido programático corresponden a 21 horas de '
    'encuentro efectivas (7 × 3 h); se registran las declaradas por coherencia con el contenido '
    'programático aprobado. Evaluación por productos de corte: primer corte 35 %, segundo corte '
    '35 %, producto final 30 %.'
)


# ── Utilidades sobre la plantilla ─────────────────────────────────────
def escribir(celda, texto):
    """Escribe en la celda respetando el formato del primer run existente."""
    parrafos = celda.paragraphs
    base = parrafos[0]
    modelo = base.runs[0] if base.runs else None

    def clonar_formato(run):
        if modelo is None:
            return
        run.font.name = modelo.font.name
        run.font.size = modelo.font.size
        run.font.bold = modelo.font.bold
        run.font.italic = modelo.font.italic
        if modelo.font.color and modelo.font.color.rgb:
            run.font.color.rgb = modelo.font.color.rgb

    for p in parrafos[1:]:
        p._element.getparent().remove(p._element)
    for r in list(base.runs):
        r._element.getparent().remove(r._element)

    lineas = texto.split('\n')
    run = base.add_run(lineas[0])
    clonar_formato(run)
    for linea in lineas[1:]:
        nuevo = copy.deepcopy(base._element)
        base._element.addnext(nuevo)
        base = docx.text.paragraph.Paragraph(nuevo, base._parent)
        for r in list(base.runs):
            r._element.getparent().remove(r._element)
        run = base.add_run(linea)
        clonar_formato(run)


def main():
    doc = docx.Document(str(PLANTILLA))
    tabla = doc.tables[0]

    # 1 · Datos de identificación
    for indice, (etiqueta, valor) in DATOS.items():
        fila = tabla.rows[indice]
        escribir(fila.cells[0], etiqueta)
        escribir(fila.cells[3], valor)

    # 2 · Contenidos. Las filas de contenido de la plantilla van de la 12
    #     en adelante; se reutilizan las que hagan falta y se borra el resto.
    primera = 12
    disponibles = len(tabla.rows) - primera
    if len(FILAS) > disponibles:
        print('ERROR: la plantilla tiene %d filas de contenido y se necesitan %d'
              % (disponibles, len(FILAS)))
        return 1

    for i, (semana, fecha, bloque, tema, responsable) in enumerate(FILAS):
        fila = tabla.rows[primera + i]
        escribir(fila.cells[0], semana)
        escribir(fila.cells[1], fecha + ('\n' + bloque if bloque else ''))
        escribir(fila.cells[2], tema)      # celdas 2 y 3 van combinadas
        escribir(fila.cells[4], responsable)

    # Borrar las filas sobrantes (PRIMER PARCIAL, ROTACIÓN, EXAMEN FINAL y
    # los renglones en blanco): este espacio evalúa por productos de corte,
    # no por parciales, y no tiene rotaciones.
    for fila in list(tabla.rows)[primera + len(FILAS):]:
        fila._tr.getparent().remove(fila._tr)

    # 3 · Nota al pie
    p = doc.add_paragraph()
    run = p.add_run(NOTA)
    run.font.size = Pt(8)
    run.font.italic = True

    doc.save(str(SALIDA))
    print('Escrito: %s' % SALIDA)
    print('  filas de contenido: %d' % len(FILAS))
    print('  filas totales en la tabla: %d' % len(tabla.rows))
    return 0


if __name__ == '__main__':
    sys.exit(main())
